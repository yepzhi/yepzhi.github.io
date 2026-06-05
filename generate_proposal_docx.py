import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_proposal_docx(filename):
    doc = docx.Document()
    
    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(30, 41, 59) # Slate 800
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("PROPUESTA DE PILOTO ESCOLAR: LABORATÓRIO DE INNOVACIÓN JÓVENESSTEM®")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(22, 51, 94) # Brand Navy
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Una Plataforma Digital Gratuita para el Fortalecimiento de la Educación Científica y Tecnológica en Secundaria/Preparatoria.")
    sub_run.italic = True
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(100, 116, 139) # Slate 500
    
    # Metadata
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run("Propuesta de Vinculación y Demo de Aula · Hermosillo, Sonora · 2026")
    meta_run.font.size = Pt(9.5)
    meta_run.font.color.rgb = RGBColor(148, 163, 184) # Slate 400
    
    # Platform Link
    link_p = doc.add_paragraph()
    link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    link_run = link_p.add_run("🌐 Plataforma oficial: https://yepzhi.com/jsweb/")
    link_run.bold = True
    link_run.font.size = Pt(11)
    link_run.font.color.rgb = RGBColor(39, 126, 255) # Brand Blue
    
    # Spacing
    doc.add_paragraph()
    
    # ── 1. PRESENTACIÓN ──
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Presentación del Proyecto")
    h1_run.font.color.rgb = RGBColor(39, 126, 255) # Brand Blue
    h1_run.font.size = Pt(14)
    h1_run.bold = True
    
    p = doc.add_paragraph(
        "JóvenesSTEM® es una iniciativa digital mexicana diseñada para democratizar el acceso al conocimiento "
        "en Ciencia, Tecnología y Programación en la educación básica y media. El programa busca incentivar el "
        "interés vocacional hacia carreras tecnológicas y de ingeniería desde la secundaria, implementando un "
        "modelo educativo interactivo donde el alumno es el protagonista de su aprendizaje."
    )
    
    # ── 2. EL PROBLEMA EN LA EDUCACIÓN SECUNDARIA ──
    h2 = doc.add_heading(level=1)
    h2_run = h2.add_run("2. El Desafío Tecnológico en las Escuelas")
    h2_run.font.color.rgb = RGBColor(39, 126, 255)
    h2_run.font.size = Pt(14)
    h2_run.bold = True
    
    p = doc.add_paragraph(
        "Las instituciones educativas de secundaria en México y LATAM enfrentan retos constantes de infraestructura. "
        "La adquisición de licencias de software, el equipamiento de laboratorios físicos de cómputo y la inestabilidad "
        "de las conexiones a internet del plantel suelen limitar o impedir el uso de herramientas educativas en la nube. "
        "JóvenesSTEM® fue desarrollado desde Sonora específicamente bajo estos lineamientos, resolviendo la barrera de "
        "infraestructura técnica mediante un entorno web ligero que aprovecha los dispositivos que los estudiantes ya "
        "tienen a la mano."
    )
    
    # ── 3. VENTAJAS PEDAGÓGICAS Y TÉCNICAS (B2B) ──
    h3 = doc.add_heading(level=1)
    h3_run = h3.add_run("3. Ventajas Clave de JóvenesSTEM® Web")
    h3_run.font.color.rgb = RGBColor(39, 126, 255)
    h3_run.font.size = Pt(14)
    h3_run.bold = True
    
    # Sub-ventajas with bold prefix
    p = doc.add_paragraph()
    p.add_run("Tutoría socrática con Inteligencia Artificial (StemBot®): ").bold = True
    p.add_run(
        "Al finalizar cada módulo de lectura, el alumno interactúa en chat con un evaluador automatizado. "
        "El bot no regala las respuestas ni realiza preguntas tradicionales de opción múltiple; en su lugar, "
        "guía al estudiante mediante cuestionamientos socráticos breves, solicitándole explicar conceptos clave "
        "con sus propias palabras para verificar la asimilación del tema."
    )
    
    p = doc.add_paragraph()
    p.add_run("Cero costo de infraestructura ($0.00 pesos): ").bold = True
    p.add_run(
        "El sistema no requiere instalaciones de software ni equipos de cómputo de última generación. "
        "Funciona perfectamente en teléfonos inteligentes (smartphones) de los estudiantes, tabletas o cualquier "
        "computadora básica con un navegador de internet (Chrome, Safari, Edge)."
    )
    
    p = doc.add_paragraph()
    p.add_run("Programa Bilingüe e Inglés A2/B1: ").bold = True
    p.add_run(
        "La plataforma está diseñada bajo un esquema completamente bilingüe (Español e Inglés). Las lecturas y "
        "actividades en inglés están redactadas bajo estándares de comprensión A2 y B1. Esto permite a los estudiantes "
        "de secundaria y preparatoria practicar el idioma técnico global de manera natural, facilitando la implementación "
        "de un modelo STEM bilingüe en el aula sin la necesidad de adquirir materiales ni licencias adicionales."
    )
    
    p = doc.add_paragraph()
    p.add_run("Monitoreo docente (Teach Portal): ").bold = True
    p.add_run(
        "El docente cuenta con un portal personalizado en tiempo real para visualizar el desempeño de su grupo. "
        "Puede observar qué alumnos han completado módulos, quiénes han aprobado la evaluación socrática "
        "y los puntos de experiencia (XP) obtenidos por el grupo."
    )
    
    p = doc.add_paragraph()
    p.add_run("Monitoreo de la Dirección/Coordinación a nivel escuela: ").bold = True
    p.add_run(
        "La plataforma incluye un panel de administración (General Manager View) donde la dirección o coordinación "
        "académica de la escuela puede supervisar el rendimiento global de todos los grupos y profesores de la "
        "institución. El panel muestra métricas consolidadas: total de alumnos activos, porcentaje de módulos "
        "completados por grupo, profesores con mayor participación y alertas de grupos con bajo rendimiento."
    )
    
    # Link to module catalog
    p_catalog = doc.add_paragraph()
    p_catalog.add_run("📚 Consulta el temario completo de 228 módulos de lectura divulgativa: ").bold = True
    link_run = p_catalog.add_run("https://yepzhi.com/jsweb/wos")
    link_run.italic = True
    link_run.font.color.rgb = RGBColor(39, 126, 255)
    
    # ── 4. PROPUESTA PILOTO 50 MINUTOS ──
    h4 = doc.add_heading(level=1)
    h4_run = h4.add_run("4. Propuesta de Piloto en el Aula (50 Minutos)")
    h4_run.font.color.rgb = RGBColor(39, 126, 255)
    h4_run.font.size = Pt(14)
    h4_run.bold = True
    
    doc.add_paragraph(
        "Para evaluar la efectividad pedagógica y técnica del sistema, proponemos llevar a cabo una sesión piloto "
        "en un grupo seleccionado de su escuela (secundaria o preparatoria). Se requiere únicamente un salón de clases "
        "con grupo de alumnos disponible, que cada alumno cuente con un teléfono celular o tableta con navegador web, "
        "y una sesión de clase de 50 minutos aproximadamente."
    )
    
    # Table for Demo Schedule
    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Style table header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Tiempo"
    hdr_cells[1].text = "Actividad"
    hdr_cells[2].text = "Detalle Técnico / Pedagógico"
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        
    # Table data
    data = [
        ("Minutos 00 - 05", "Introducción y Contexto", "Explicación breve al grupo sobre los fundamentos STEM y la lógica de programación."),
        ("Minutos 05 - 10", "Registro e Inicio", "Los estudiantes acceden a yepzhi.com/jsweb/ y crean su cuenta en un solo clic usando Google."),
        ("Minutos 10 - 20", "Lectura Interactiva", "El alumno lee el módulo de prueba asignado en la plataforma (ej. Lógica de Algoritmos)."),
        ("Minutos 20 - 45", "Acreditación Socrática", "Chat activo con StemBot. Los alumnos defienden sus argumentos para ganar su primer insignia."),
        ("Minutos 45 - 50", "Retroalimentación Docente", "El profesor proyecta el Teach Portal mostrando las estadísticas y logros del grupo en vivo.")
    ]
    
    for i, row_data in enumerate(data):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]
        
        # Format text inside table
        for idx in range(3):
            row_cells[idx].paragraphs[0].runs[0].font.size = Pt(9.5)
            
    doc.add_paragraph()
    
    # ── 5. CURRÍCULUM Y LICENCIAMIENTO ──
    h5 = doc.add_heading(level=1)
    h5_run = h5.add_run("5. Currículum y Gratuidad")
    h5_run.font.color.rgb = RGBColor(39, 126, 255)
    h5_run.font.size = Pt(14)
    h5_run.bold = True
    
    doc.add_paragraph(
        "El programa completo cuenta con 228 módulos de estudio estructurados y alineados a estándares internacionales "
        "de educación científica (Next Generation Science Standards - NGSS). Todo el acceso al portal de estudiantes, "
        "la interacción de tutoría inteligente de StemBot y las herramientas de administración docente se proporcionan "
        "con licenciamiento abierto y gratuito para las escuelas secundarias y preparatorias públicas y privadas. "
        "El proyecto no tiene fines lucrativos comerciales y busca crear una comunidad de aprendizaje digital colaborativa. "
        "Solamente y en el caso opcional que el alumno desee obtener su certificado \"JóvenesSTEM\" tendrá un costo "
        "de $49 MXN pesos para mantenimiento de servidores de IA, pero el conocimiento tecnológico y científico es "
        "democratizado y aseguramos el acceso universal de forma gratuita con esta iniciativa."
    )
    
    # Horizontal line separator
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep.add_run("────────────────────────────────────────────────────────────").font.color.rgb = RGBColor(200, 200, 200)
    
    # Contact information
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_contact.add_run("JóvenesSTEM® Web · Vinculación Escolar e Innovación EdTech\n")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(22, 51, 94)
    
    r2 = p_contact.add_run("Contacto y Agenda de Pruebas: Alberto Yépiz (Fundador) · Correo: yepzhi@gmail.com\n")
    r2.font.size = Pt(9.5)
    
    r3 = p_contact.add_run("Plataforma oficial: https://yepzhi.com/jsweb/\n")
    r3.italic = True
    r3.font.size = Pt(9.5)
    r3.font.color.rgb = RGBColor(39, 126, 255)
    
    r4 = p_contact.add_run("Catálogo de módulos: https://yepzhi.com/jsweb/wos")
    r4.italic = True
    r4.font.size = Pt(9.5)
    r4.font.color.rgb = RGBColor(0, 168, 150)

    doc.save(filename)
    print(f"✅ Generated {filename} successfully.")

if __name__ == "__main__":
    create_proposal_docx("propuesta_piloto_jovenesstem_secundaria.docx")
